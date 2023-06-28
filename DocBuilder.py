from tkinter import messagebox
import pythoncom
import tkinter as tk
from docx import Document
import xlwings as xw
import os
from win32com.client import Dispatch
import threading

class DocBuilder():
    def __init__(self) -> None:
        
        self.root = tk.Tk()
        self.root.title("DocBuilder")
        self.height = 250
        self.width = 250
        self.x = self.root.winfo_screenwidth() / 2 - self.width / 2
        self.y = self.root.winfo_screenheight() / 2 - self.height / 2
        self.root.geometry('%dx%d+%d+%d' % (self.width, self.height, self.x, self.y))
        
        self.lbl_Author = tk.Label(self.root, text="DocBuilder by boyce",padx=5,pady=5)
        self.lbl_Author.pack(padx=3,pady=3)
        
        self.lbl_File = tk.LabelFrame(self.root, text="你想生成的文件类型?",padx=5,pady=5)
        self.lbl_File.pack(padx=10,pady=10)
        

        
        self.varWord = tk.IntVar()
        
        word_btn = tk.Checkbutton(self.lbl_File, text="word", variable=self.varWord, onvalue= 888, offvalue=0).pack(anchor='w')
        
        self.varPdf = tk.IntVar()

        pdf_btn = tk.Checkbutton(self.lbl_File, text="pdf", variable=self.varPdf, onvalue= 1, offvalue=0).pack(anchor='w')
        
        self.varExcel = tk.IntVar()
        
        excel_btn = tk.Checkbutton(self.lbl_File, text="excel", variable=self.varExcel, onvalue= 999, offvalue=0).pack(anchor='w')
        
        self.lbl_FileCount = tk.LabelFrame(self.root, text="你想生成的文件数量?",padx=5,pady=5)
        self.lbl_FileCount.pack(padx=10,pady=10)
        
        self.varCount = tk.IntVar()
        
        count_entry = tk.Entry(self.lbl_FileCount, width=8, textvariable= self.varCount).pack(side='left')
        
        self.btn_create = tk.Button(self.lbl_FileCount, text="我要创建", command=self.__CreateFunc).pack(side='left')

        self.btn_delete = tk.Button(self.lbl_FileCount, text="删除已创建", command=self.__DeleteAll).pack(side='left')

        
        self.root.mainloop()

        
    def __CreateFunc(self):
        # print(self.varWord.get())
        # print(self.varPdf.get())
        # print(self.varExcel.get())
        # print(self.varCount.get())
        
        if self.varCount.get() == 0 or (self.varWord.get() == 0 and self.varPdf.get() == 0 and self.varExcel.get() == 0):
            messagebox.showerror(title="错误", message="文件类型未选择或数量不能为0")
        else:
            if self.varWord.get() == 888:
                threading.Thread(target=self.__funCreateWord, args=(self.varCount.get(),)).start()
                # print("Create Word")
                
            if self.varPdf.get()== 1:
                threading.Thread(target=self.__funCreatePdf, args=(self.varCount.get(),)).start()

                # print("Create Pdf")
                
            if self.varExcel.get()== 999:
                threading.Thread(target=self.__funCreateExcel, args=(self.varCount.get(),)).start()

                # print("Create Excel")
                
    def __DeleteAll(self):
        
        for i in os.listdir():
            if i.endswith("_test.docx") or i.endswith("_test.pdf") or i.endswith("_test.xlsx"):
                os.remove(i)
        messagebox.showinfo(title="成功", message="目录在:\n"+os.getcwd()+"\n的测试文件已全部删除")
    # 创建新的word
    def __funCreateWord(self, num):
        if num == 0:
            pass
        else:
            # pythoncom.CoInitialize()
            for i in range(1, num + 1):
                document = Document()
                document.add_paragraph(f'para{i}: This is a new paragraph. ')
                document.save(rf'Word{i}_test.docx')
            # pythoncom.CoUninitialize()
        messagebox.showinfo(title="成功", message="word文件已成功生成，目录在:\n"+os.getcwd()+"\n")
            
    # 创建新的pdf
    def __funCreatePdf(self, num):
        if num == 0:
            pass
        else:
            pythoncom.CoInitialize()
            word = Dispatch("word.application")
            word.visible = 0
            word.DisplayAlerts = 0
            for i in range(1, num + 1):
                document = Document()
                document.add_paragraph(f'para{i}: This is a new paragraph. ')
                document.save(rf'w{i}_test.docx')
                pdf = word.Documents.Open(rf'{os.getcwd()}\w{i}_test.docx')
                pdf.SaveAs(rf'{os.getcwd()}\Pdf{i}_test.pdf', 17)
                pdf.Close()
                os.remove(rf'{os.getcwd()}\w{i}_test.docx')
            word.Quit()
            pythoncom.CoUninitialize()
        messagebox.showinfo(title="成功", message="pdf文件已成功生成，目录在:\n"+os.getcwd()+"\n")
    
    # 创建新的excel       
    def __funCreateExcel(self, num):
        if num == 0:
            pass
        else:
            app = xw.App(visible=False, add_book=False)

            for i in range(1, num + 1):
                wordbook = app.books.add()
                sheet = wordbook.sheets['Sheet1']
                sheet[f'A{i}'].value = f'cell {i}'
                wordbook.save(f'Excel{i}_test.xlsx')
                wordbook.close()

            app.quit()
        messagebox.showinfo(title="成功", message="excel文件已成功生成，目录在:\n"+os.getcwd()+"\n")

        
if __name__ == '__main__':
    
    print("当前文件路径：",os.getcwd())
    DocBuilder()
    # print("调用funCreateWord()")
    # funCreateWord(3,1)
    
    # print("调用funCreateExcel()")
    # funCreateExcel(5)